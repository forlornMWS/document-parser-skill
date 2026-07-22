#!/usr/bin/env python3
"""
Word 模板格式解析器
将 docx 文件解析为结构化 JSON，供 Agent 生成 PDF 时参考样式。

用法:
    python parse_word.py <input.docx> [output.json]
    python parse_word.py template.docx                    # 输出到 template_format.json
    python parse_word.py template.docx -                  # 输出到 stdout
    python parse_word.py template.docx --pretty           # 美化输出(带缩进)
    python parse_word.py template.docx --skip-empty       # 跳过空段落
    python parse_word.py template.docx --extract-images   # 提取图片到 images/ 目录
"""

import sys
import json
import argparse
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.text.run import Run
    from docx.text.paragraph import Paragraph
    from docx.table import Table
except ImportError:
    print("Error: python-docx or lxml not installed. Run: pip install python-docx lxml", file=sys.stderr)
    sys.exit(1)


# ── 常量映射 ──────────────────────────────────────────────

PAPER_SIZES = {
    (210, 297): "A4",
    (297, 210): "A4_landscape",
    (216, 279): "Letter",
    (279, 216): "Letter_landscape",
    (216, 356): "Legal",
    (356, 216): "Legal_landscape",
    (297, 420): "A3",
    (420, 297): "A3_landscape",
    (148, 210): "A5",
    (210, 148): "A5_landscape",
}

ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
    None: "left",
}

LINE_SPACING_MAP = {
    WD_LINE_SPACING.SINGLE: "single",
    WD_LINE_SPACING.ONE_POINT_FIVE: "1.5",
    WD_LINE_SPACING.DOUBLE: "double",
    WD_LINE_SPACING.AT_LEAST: "at_least",
    WD_LINE_SPACING.EXACTLY: "exactly",
    WD_LINE_SPACING.MULTIPLE: "multiple",
    None: "single",
}

TABLE_ALIGNMENT_MAP = {
    WD_TABLE_ALIGNMENT.LEFT: "left",
    WD_TABLE_ALIGNMENT.CENTER: "center",
    WD_TABLE_ALIGNMENT.RIGHT: "right",
    None: "left",
}

CELL_VERTICAL_ALIGN_MAP = {
    WD_CELL_VERTICAL_ALIGNMENT.TOP: "top",
    WD_CELL_VERTICAL_ALIGNMENT.CENTER: "center",
    WD_CELL_VERTICAL_ALIGNMENT.BOTTOM: "bottom",
    WD_CELL_VERTICAL_ALIGNMENT.BOTH: "justify",
    None: "top",
}


def emu_to_pt(emu):
    """EMU 转磅"""
    if emu is None:
        return None
    return round(emu / 12700, 2)


def emu_to_cm(emu):
    """EMU 转厘米"""
    if emu is None:
        return None
    return round(emu / 360000, 2)


def emu_to_px(emu):
    """EMU 转像素 (96 DPI)"""
    if emu is None:
        return None
    return round(emu / 914400 * 96, 1)


def parse_border_element(border_elem):
    """解析单个边框元素的通用函数（段落/单元格/表格通用）"""
    border_info = {}
    val = border_elem.get(qn('w:val'))
    if val:
        border_info['style'] = val
    sz = border_elem.get(qn('w:sz'))
    if sz:
        border_info['size'] = int(sz) / 8  # 转为磅
    space = border_elem.get(qn('w:space'))
    if space:
        border_info['space_pt'] = int(space)
    color = border_elem.get(qn('w:color'))
    if color:
        border_info['color'] = color
    return border_info


def parse_color(color_obj):
    """解析颜色对象"""
    if color_obj is None:
        return None
    
    # RGB 颜色
    if hasattr(color_obj, 'rgb') and color_obj.rgb:
        return str(color_obj.rgb)
    
    # 主题颜色
    if hasattr(color_obj, 'theme') and color_obj.theme is not None:
        theme_val = color_obj.theme
        # MSO_THEME_COLOR 枚举，取 name 避免输出 "MSO_THEME_COLOR.TEXT_1 (0)"
        theme_name = getattr(theme_val, 'name', None) or str(theme_val)
        return f"theme:{theme_name}"
    
    return None


def parse_font(run):
    """解析字体样式"""
    font = run.font
    result = {}
    
    if font.name:
        result['name'] = font.name
    
    # 东亚字体
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is not None:
            east_asia = rfonts.get(qn('w:eastAsia'))
            if east_asia:
                result['east_asia'] = east_asia
    
    if font.size:
        result['size_pt'] = font.size.pt
    
    if font.bold:
        result['bold'] = True
    
    if font.italic:
        result['italic'] = True
    
    if font.underline:
        result['underline'] = True if font.underline is True else str(font.underline)
    
    if font.strike:
        result['strike'] = True
    
    if font.all_caps:
        result['all_caps'] = True
    
    if font.small_caps:
        result['small_caps'] = True
    
    # 颜色
    color = parse_color(font.color)
    if color:
        result['color'] = color
    
    # 高亮
    if font.highlight_color:
        result['highlight'] = str(font.highlight_color)
    
    # 上标/下标
    if font.superscript:
        result['superscript'] = True
    if font.subscript:
        result['subscript'] = True
    
    return result if result else None


def parse_paragraph_format(paragraph):
    """解析段落格式"""
    pf = paragraph.paragraph_format
    result = {}
    
    # 对齐方式
    if pf.alignment is not None:
        result['alignment'] = ALIGN_MAP.get(pf.alignment, 'left')
    
    # 缩进
    if pf.left_indent:
        result['left_indent_pt'] = emu_to_pt(pf.left_indent)
    if pf.right_indent:
        result['right_indent_pt'] = emu_to_pt(pf.right_indent)
    if pf.first_line_indent:
        result['first_line_indent_pt'] = emu_to_pt(pf.first_line_indent)
    
    # 行距
    if pf.line_spacing_rule is not None:
        result['line_spacing_rule'] = LINE_SPACING_MAP.get(pf.line_spacing_rule, 'single')
    if pf.line_spacing:
        # 如果是 exactly 或 at_least，line_spacing 是 EMU 对象，需要转磅
        if pf.line_spacing_rule in (WD_LINE_SPACING.EXACTLY, WD_LINE_SPACING.AT_LEAST):
            result['line_spacing_pt'] = emu_to_pt(pf.line_spacing)
        elif isinstance(pf.line_spacing, (int, float)):
            result['line_spacing'] = pf.line_spacing
        else:
            result['line_spacing_pt'] = emu_to_pt(pf.line_spacing)
    
    # 段前段后间距
    if pf.space_before:
        result['space_before_pt'] = emu_to_pt(pf.space_before)
    if pf.space_after:
        result['space_after_pt'] = emu_to_pt(pf.space_after)
    
    # 分页
    if pf.page_break_before:
        result['page_break_before'] = True
    
    # 保持行/段
    if pf.keep_together:
        result['keep_together'] = True
    if pf.keep_with_next:
        result['keep_with_next'] = True
    
    # 大纲级别和段落边框（通过 XML 获取）
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is not None:
            val = outlineLvl.get(qn('w:val'))
            if val is not None:
                result['outline_level'] = int(val)
        
        # 段落边框（如红色横线、四边边框等）
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            borders = {}
            for side_name in ['top', 'left', 'bottom', 'right', 'between', 'bar']:
                side_elem = pBdr.find(qn(f'w:{side_name}'))
                if side_elem is not None:
                    border_info = parse_border_element(side_elem)
                    if border_info:
                        borders[side_name] = border_info
            if borders:
                result['borders'] = borders
    
    return result if result else None


def parse_run(run):
    """解析文本运行"""
    result = {
        'text': run.text,
    }
    
    font = parse_font(run)
    if font:
        result['font'] = font
    
    return result


def parse_drawingml_shape(drawing_elem):
    """解析 DrawingML 形状（线条、矩形等）"""
    shapes = []
    
    # 定义命名空间
    WPS_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    
    # 查找 wsp:wsp 元素（WordprocessingShape）
    wsp_elements = drawing_elem.findall('.//{%s}wsp' % WPS_NS)
    
    for wsp in wsp_elements:
        shape_info = {
            'type': 'drawingml_shape'
        }
        
        # 获取形状名称
        cnv_sp_pr = wsp.find('.//{%s}cNvSpPr' % WPS_NS)
        if cnv_sp_pr is not None:
            shape_info['name'] = cnv_sp_pr.get('name', '')
        
        # 获取形状属性
        sp_pr = wsp.find('.//{%s}spPr' % WPS_NS)
        if sp_pr is not None:
            # 获取预设几何形状（如 line, rect 等）
            prst_geom = sp_pr.find('{%s}prstGeom' % A_NS)
            if prst_geom is not None:
                shape_info['preset'] = prst_geom.get('prst', '')
            
            # 获取线条样式
            ln = sp_pr.find('{%s}ln' % A_NS)
            if ln is not None:
                line_info = {}
                # 线条宽度（EMU 转磅）
                w = ln.get('w')
                if w:
                    line_info['width_pt'] = round(int(w) / 12700, 2)
                
                # 线条颜色
                solid_fill = ln.find('{%s}solidFill' % A_NS)
                if solid_fill is not None:
                    srgb_clr = solid_fill.find('{%s}srgbClr' % A_NS)
                    if srgb_clr is not None:
                        line_info['color'] = srgb_clr.get('val', '')
                
                if line_info:
                    shape_info['line'] = line_info
            
            # 获取尺寸和位置
            xfrm = sp_pr.find('{%s}xfrm' % A_NS)
            if xfrm is not None:
                ext = xfrm.find('{%s}ext' % A_NS)
                if ext is not None:
                    cx = ext.get('cx')
                    cy = ext.get('cy')
                    if cx:
                        shape_info['width_emu'] = int(cx)
                        shape_info['width_pt'] = round(int(cx) / 12700, 2)
                    if cy:
                        shape_info['height_emu'] = int(cy)
                        shape_info['height_pt'] = round(int(cy) / 12700, 2)
                
                off = xfrm.find('{%s}off' % A_NS)
                if off is not None:
                    x = off.get('x')
                    y = off.get('y')
                    if x:
                        shape_info['offset_x_emu'] = int(x)
                    if y:
                        shape_info['offset_y_emu'] = int(y)
        
        shapes.append(shape_info)
    
    return shapes


def parse_vml_shape(pict_elem):
    """解析 VML 形状（旧格式兼容）"""
    shapes = []
    
    # 定义命名空间
    VML_NS = 'urn:schemas-microsoft-com:vml'
    
    # 查找 v:shape 元素
    v_shape_elements = pict_elem.findall('.//{%s}shape' % VML_NS)
    
    for v_shape in v_shape_elements:
        shape_info = {
            'type': 'vml_shape'
        }
        
        # 获取形状 ID
        shape_info['id'] = v_shape.get('id', '')
        
        # 获取样式
        style = v_shape.get('style', '')
        if style:
            shape_info['style'] = style
        
        # 获取线条样式
        stroke = v_shape.find('{%s}stroke' % VML_NS)
        if stroke is not None:
            stroke_info = {}
            color = stroke.get('color')
            if color:
                stroke_info['color'] = color
            weight = stroke.get('weight')
            if weight:
                stroke_info['weight'] = weight
            if stroke_info:
                shape_info['stroke'] = stroke_info
        
        # 获取填充
        fill = v_shape.find('{%s}fill' % VML_NS)
        if fill is not None:
            fill_info = {}
            color = fill.get('color')
            if color:
                fill_info['color'] = color
            if fill_info:
                shape_info['fill'] = fill_info
        
        shapes.append(shape_info)
    
    return shapes


def parse_paragraph(paragraph, extract_images=False, image_dir=None, image_counter=None):
    """解析段落"""
    result = {
        'type': 'paragraph',
        'text': paragraph.text,
    }
    
    # 段落样式
    if paragraph.style:
        result['style_name'] = paragraph.style.name
    
    # 段落格式
    fmt = parse_paragraph_format(paragraph)
    if fmt:
        result['format'] = fmt
    
    # 列表样式
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            numId = numPr.find(qn('w:numId'))
            if ilvl is not None and numId is not None:
                result['list_level'] = int(ilvl.get(qn('w:val'), '0'))
                result['list_id'] = int(numId.get(qn('w:val'), '0'))
    
    # 文本运行
    runs = []
    shapes = []
    
    for child in paragraph._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        
        if tag == 'r':
            # 普通 run
            run = Run(child, paragraph)
            run_data = parse_run(run)
            
            # 检查是否有绘图元素
            drawings = child.findall(qn('w:drawing'))
            for drawing in drawings:
                # 解析 DrawingML 形状
                drawingml_shapes = parse_drawingml_shape(drawing)
                if drawingml_shapes:
                    shapes.extend(drawingml_shapes)
                
                # 提取图片（如果启用）
                if extract_images and image_dir and image_counter is not None:
                    blip = drawing.find('.//' + qn('a:blip'))
                    if blip is not None:
                        embed_id = blip.get(qn('r:embed'))
                        if embed_id:
                            # 从文档关系中获取图片
                            try:
                                image_part = run.part.related_parts.get(embed_id)
                                if image_part and hasattr(image_part, 'image'):
                                    image = image_part.image
                                    image_ext = image.content_type.split('/')[-1]
                                    if image_ext == 'jpeg':
                                        image_ext = 'jpg'
                                    image_filename = f"image_{image_counter[0]}.{image_ext}"
                                    image_path = image_dir / image_filename
                                    with open(image_path, 'wb') as f:
                                        f.write(image.blob)
                                    run_data['image'] = {
                                        'filename': image_filename,
                                        'content_type': image.content_type,
                                        'size_bytes': len(image.blob)
                                    }
                                    image_counter[0] += 1
                            except Exception as e:
                                print(f"Warning: Failed to extract image: {e}", file=sys.stderr)
            
            runs.append(run_data)
        
        elif tag == 'pict':
            # VML 形状（旧格式）
            vml_shapes = parse_vml_shape(child)
            if vml_shapes:
                shapes.extend(vml_shapes)
        
        elif tag == 'hyperlink':
            # 超链接
            hyperlink_text = ''
            hyperlink_url = None
            hyperlink_runs = []
            
            # 获取超链接 URL
            r_id = child.get(qn('r:id'))
            if r_id:
                try:
                    rel = paragraph.part.rels[r_id]
                    hyperlink_url = rel.target_ref
                except Exception:
                    pass
            
            # 解析超链接内的 runs
            for r_elem in child.findall(qn('w:r')):
                run = Run(r_elem, paragraph)
                run_data = parse_run(run)
                hyperlink_text += run.text
                hyperlink_runs.append(run_data)
            
            if hyperlink_url:
                hyperlink_data = {
                    'text': hyperlink_text,
                    'hyperlink': hyperlink_url,
                }
                if hyperlink_runs:
                    hyperlink_data['runs'] = hyperlink_runs
                runs.append(hyperlink_data)
    
    if runs:
        result['runs'] = runs
    
    if shapes:
        result['shapes'] = shapes
    
    return result


def parse_cell_border(cell):
    """解析单元格边框"""
    tc = cell._tc
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is None:
        return None
    
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        return None
    
    result = {}
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_elem = tc_borders.find(qn(f'w:{border_name}'))
        if border_elem is not None:
            border_info = parse_border_element(border_elem)
            if border_info:
                result[border_name] = border_info
    
    return result if result else None


def parse_cell(cell, extract_images=False, image_dir=None, image_counter=None):
    """解析表格单元格"""
    result = {
        'text': cell.text,
    }
    
    # 垂直对齐
    if cell.vertical_alignment is not None:
        result['vertical_alignment'] = CELL_VERTICAL_ALIGN_MAP.get(cell.vertical_alignment, 'top')
    
    # 背景色
    shading = cell._tc.find(qn('w:tcPr') + '/' + qn('w:shd'))
    if shading is not None:
        fill = shading.get(qn('w:fill'))
        if fill and fill != 'auto':
            result['background_color'] = fill
    
    # 边框
    borders = parse_cell_border(cell)
    if borders:
        result['borders'] = borders
    
    # 合并信息
    tc = cell._tc
    grid_span = tc.find(qn('w:tcPr') + '/' + qn('w:gridSpan'))
    if grid_span is not None:
        val = grid_span.get(qn('w:val'))
        if val:
            result['colspan'] = int(val)
    
    v_merge = tc.find(qn('w:tcPr') + '/' + qn('w:vMerge'))
    if v_merge is not None:
        val = v_merge.get(qn('w:val'))
        if val == 'restart':
            result['vmerge_restart'] = True
        else:
            result['vmerge_continue'] = True
    
    # 段落
    paragraphs = []
    for para in cell.paragraphs:
        paragraphs.append(parse_paragraph(para, extract_images, image_dir, image_counter))
    if paragraphs:
        result['paragraphs'] = paragraphs
    
    return result


def parse_table(table, extract_images=False, image_dir=None, image_counter=None):
    """解析表格"""
    result = {
        'type': 'table',
    }
    
    # 表格对齐
    if table.alignment is not None:
        result['alignment'] = TABLE_ALIGNMENT_MAP.get(table.alignment, 'left')
    
    # 表格样式
    if table.style:
        result['style_name'] = table.style.name
    
    # 表格属性
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is not None:
        # 表格边框
        tbl_borders = tbl_pr.find(qn('w:tblBorders'))
        if tbl_borders is not None:
            borders = {}
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_elem = tbl_borders.find(qn(f'w:{border_name}'))
                if border_elem is not None:
                    border_info = parse_border_element(border_elem)
                    if border_info:
                        borders[border_name] = border_info
            if borders:
                result['borders'] = borders
        
        # 表格宽度
        tbl_w = tbl_pr.find(qn('w:tblW'))
        if tbl_w is not None:
            w = tbl_w.get(qn('w:w'))
            w_type = tbl_w.get(qn('w:type'))
            if w:
                result['width'] = int(w)
                result['width_type'] = w_type or 'auto'
    
    # 行和单元格
    rows = []
    for row in table.rows:
        row_data = {
            'cells': []
        }
        
        # 行高
        tr = row._tr
        tr_pr = tr.find(qn('w:trPr'))
        if tr_pr is not None:
            tr_height = tr_pr.find(qn('w:trHeight'))
            if tr_height is not None:
                val = tr_height.get(qn('w:val'))
                rule = tr_height.get(qn('w:hRule'))
                if val:
                    row_data['height_pt'] = int(val) / 20  # twip 转磅
                    row_data['height_rule'] = rule or 'at_least'
        
        for cell in row.cells:
            row_data['cells'].append(parse_cell(cell, extract_images, image_dir, image_counter))
        
        rows.append(row_data)
    
    result['rows'] = rows
    
    return result


def parse_section(section):
    """解析节（页面设置）"""
    result = {}
    
    # 纸张大小
    width_cm = emu_to_cm(section.page_width)
    height_cm = emu_to_cm(section.page_height)
    
    if width_cm and height_cm:
        result['page_width_cm'] = width_cm
        result['page_height_cm'] = height_cm
        
        # 尝试识别标准纸张
        width_mm = int(width_cm * 10)
        height_mm = int(height_cm * 10)
        paper_key = (width_mm, height_mm)
        result['paper_size'] = PAPER_SIZES.get(paper_key, f"Custom_{width_mm}x{height_mm}mm")
    
    # 页面方向
    if section.orientation is not None:
        result['orientation'] = 'landscape' if section.orientation == WD_ORIENT.LANDSCAPE else 'portrait'
    
    # 页边距
    result['margins'] = {
        'top_cm': emu_to_cm(section.top_margin),
        'bottom_cm': emu_to_cm(section.bottom_margin),
        'left_cm': emu_to_cm(section.left_margin),
        'right_cm': emu_to_cm(section.right_margin),
        'header_cm': emu_to_cm(section.header_distance),
        'footer_cm': emu_to_cm(section.footer_distance),
    }
    
    # 页眉页脚
    if section.different_first_page_header_footer:
        result['different_first_page'] = True
    
    return result


def parse_header_footer(header_or_footer):
    """解析页眉或页脚"""
    if not header_or_footer:
        return None
    if header_or_footer.is_linked_to_previous:
        return {'linked_to_previous': True}
    paragraphs = [parse_paragraph(para) for para in header_or_footer.paragraphs]
    return {'linked_to_previous': False, 'paragraphs': paragraphs} if paragraphs else None


def has_shapes(paragraph):
    """检查段落是否包含形状（DrawingML 或 VML）"""
    # 定义命名空间
    WPS_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    VML_NS = 'urn:schemas-microsoft-com:vml'
    
    # 检查 w:drawing 中的 wps:wsp（DrawingML 形状）
    drawings = paragraph._element.findall('.//' + qn('w:drawing'))
    for drawing in drawings:
        wsp_elements = drawing.findall('.//{%s}wsp' % WPS_NS)
        if wsp_elements:
            return True
    
    # 检查 w:pict 中的 v:shape（VML 形状）
    picts = paragraph._element.findall('.//' + qn('w:pict'))
    for pict in picts:
        v_shapes = pict.findall('.//{%s}shape' % VML_NS)
        if v_shapes:
            return True
    
    return False


def parse_document(filepath, extract_images=False, skip_empty=False):
    """解析整个文档"""
    doc = Document(filepath)
    
    result = {
        'metadata': {
            'source_file': Path(filepath).name,
            'parsed_at': datetime.now().isoformat(),
            'parser_version': '1.0.0',
        },
        'sections': [],
        'body': [],
    }
    
    # 解析节（页面设置）
    for section in doc.sections:
        result['sections'].append(parse_section(section))
    
    # 准备图片提取
    image_dir = None
    image_counter = [0]
    if extract_images:
        image_dir = Path(filepath).parent / f"{Path(filepath).stem}_images"
        image_dir.mkdir(exist_ok=True)
        result['metadata']['image_dir'] = str(image_dir)
    
    # 解析正文内容
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        
        if tag == 'p':
            # 段落
            para = Paragraph(element, doc)
            # 修改：不仅检查文本，还要检查是否包含形状
            if skip_empty and not para.text.strip() and not has_shapes(para):
                continue
            result['body'].append(parse_paragraph(para, extract_images, image_dir, image_counter))
        
        elif tag == 'tbl':
            # 表格
            table = Table(element, doc)
            result['body'].append(parse_table(table, extract_images, image_dir, image_counter))
        
        elif tag == 'sectPr':
            # 节属性（已在 sections 中处理）
            pass
    
    # 解析页眉页脚（取第一个节的）
    if doc.sections:
        section = doc.sections[0]
        
        # 默认页眉
        header = parse_header_footer(section.header)
        if header:
            result['header'] = header
        
        # 默认页脚
        footer = parse_header_footer(section.footer)
        if footer:
            result['footer'] = footer
        
        # 首页页眉页脚
        if section.different_first_page_header_footer:
            first_header = parse_header_footer(section.first_page_header)
            if first_header:
                result['first_page_header'] = first_header
            
            first_footer = parse_header_footer(section.first_page_footer)
            if first_footer:
                result['first_page_footer'] = first_footer
    
    # 统计
    para_count = sum(1 for item in result['body'] if item.get('type') == 'paragraph')
    table_count = sum(1 for item in result['body'] if item.get('type') == 'table')
    result['metadata']['statistics'] = {
        'paragraphs': para_count,
        'tables': table_count,
        'sections': len(result['sections']),
    }
    
    if extract_images:
        result['metadata']['statistics']['images'] = image_counter[0]
    
    return result


def main():
    parser = argparse.ArgumentParser(
        description="Word 模板格式解析器 - 将 docx 转为结构化 JSON"
    )
    parser.add_argument('input', help='输入的 docx 文件路径')
    parser.add_argument('output', nargs='?', default=None,
                        help='输出 JSON 文件路径 (默认: <input>_format.json, 用 - 输出到 stdout)')
    parser.add_argument('--pretty', action='store_true',
                        help='美化输出(带缩进，默认紧凑)')
    parser.add_argument('--skip-empty', action='store_true',
                        help='跳过空段落')
    parser.add_argument('--extract-images', action='store_true',
                        help='提取图片到 images/ 目录')
    
    args = parser.parse_args()
    
    # 检查输入文件
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)
    
    # 解析
    print(f"Parsing: {args.input}", file=sys.stderr)
    data = parse_document(args.input, extract_images=args.extract_images, skip_empty=args.skip_empty)
    
    # 统计
    stats = data['metadata']['statistics']
    print(f"Sections: {stats['sections']}, Paragraphs: {stats['paragraphs']}, Tables: {stats['tables']}", file=sys.stderr)
    if args.extract_images:
        print(f"Images extracted: {stats['images']}", file=sys.stderr)
    
    # 输出
    indent = 2 if args.pretty else None
    json_str = json.dumps(data, indent=indent, ensure_ascii=False, default=str)
    
    if args.output == '-':
        sys.stdout.write(json_str)
        sys.stdout.write("\n")
    else:
        output_path = args.output or str(input_path.with_stem(input_path.stem + "_format"))
        if not output_path.endswith('.json'):
            output_path += '.json'
        Path(output_path).write_text(json_str, encoding='utf-8')
        print(f"Output: {output_path}", file=sys.stderr)


if __name__ == '__main__':
    main()
